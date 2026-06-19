"""
FastAPI backend for SkillSync AI.

Local dev:
    pip install -r requirements.txt
    export GROQ_API_KEY=your_key_here
    uvicorn api.index:app --reload --port 8000

Deployed on Vercel, this file is auto-detected as the Python entrypoint
(api/index.py) and served at /api/*.
"""
import os
import json
import tempfile
from pathlib import Path

import docx
from pypdf import PdfReader
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

DEFAULT_MODEL = "llama-3.3-70b-versatile"
DEFAULT_TEMPERATURE = 0.2

app = FastAPI(title="SkillSync AI API")

# Wide open CORS for simplicity. If the frontend and API ever live on
# different domains, tighten this to the actual frontend origin.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# File parsing
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")
    return text


@app.post("/api/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    suffix = Path(file.filename or "").suffix.lower()
    raw = await file.read()

    # Use a unique temp file per request — safe under concurrent requests.
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(raw)
        tmp_path = tmp.name

    try:
        if suffix == ".pdf":
            text = extract_text_from_pdf(tmp_path)
        elif suffix == ".docx":
            text = extract_text_from_docx(tmp_path)
        else:
            text = raw.decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        os.unlink(tmp_path)

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Couldn't find any readable text in that file. Try pasting it instead.",
        )

    return {"text": text}


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------

class AnalyzeRequest(BaseModel):
    resume_text: str
    job_desc_text: str
    target_role: str | None = None
    experience_level: str | None = None
    focus_area: str | None = None


def build_application_context(target_role=None, experience_level=None, focus_area=None) -> str:
    context_lines = []
    if target_role and target_role.strip():
        context_lines.append(f"Target role: {target_role.strip()}")
    if experience_level and experience_level != "Not specified":
        context_lines.append(f"Experience level: {experience_level}")
    if focus_area:
        context_lines.append(f"Main focus requested by user: {focus_area}")
    return "\n".join(context_lines) if context_lines else "No extra user preferences provided."


SYSTEM_PROMPT = (
    "You are an expert Applicant Tracking System (ATS) reviewer, HR consultant, and resume coach. "
    "Your task is to perform a highly thorough, accurate, and professional ATS audit of the user's resume against "
    "the provided job description. You must output the analysis strictly in JSON format. Do not write any conversational "
    "text before or after the JSON payload. Ensure all JSON fields are populated accurately. "
    "Keep the wording practical and applicant-friendly. Do not mention model names, API providers, temperature, prompts, "
    "JSON schema details, or backend implementation in any user-facing response field."
)

OUTPUT_SCHEMA = """
--- OUTPUT JSON SCHEMA ---
You must return a JSON object with the following exact keys and types:
{
  "ats_score": <int, 0 to 100 representing general compatibility>,
  "match_percentage": <int, 0 to 100 representing job description keyword/requirement alignment>,
  "formatting_score": <int, 0 to 100 representing layout/font/readability elements>,
  "keyword_score": <int, 0 to 100 representing presence of core resume keywords>,
  "experience_score": <int, 0 to 100 representing relevance of past roles to requirements>,
  "skills_score": <int, 0 to 100 representing job-related and soft skills alignment>,
  "analysis_summary": "<string, professional summary of candidate fit, highlighting core alignments and main gaps>",
  "key_strengths": ["<string>", "<string>", ...],
  "critical_issues": ["<string>", "<string>", ...],
  "keyword_analysis": [
    {"keyword": "<string>", "status": "<'Present' or 'Missing'>", "importance": "<'High'|'Medium'|'Low'>"}
  ],
  "skills_gap": [
    {"skill": "<string>", "category": "<'Hard Skill'|'Soft Skill'|'Tool/Technology'>", "gap_description": "<string>"}
  ],
  "ats_compatibility_checks": [
    {"check": "<string>", "status": "<'Passed'|'Warning'|'Failed'>", "feedback": "<string>"}
  ],
  "improvement_suggestions": ["<string>", ...],
  "optimized_resume_snippets": [
    {"original": "<string>", "optimized": "<string>", "rationale": "<string>"}
  ],
  "interview_questions": [
    {"question": "<string>", "rationale": "<string>"}
  ],
  "cover_letter": "<string, a full cover letter using the candidate's strengths>"
}
"""


def analyze_resume_ats(resume_text, job_desc_text, target_role=None, experience_level=None, focus_area=None):
    key = os.getenv("GROQ_API_KEY")
    if not key:
        raise HTTPException(status_code=503, detail="The analysis service is not configured yet.")

    client = Groq(api_key=key)
    application_context = build_application_context(target_role, experience_level, focus_area)

    user_prompt = f"""
Analyze the resume and job description below.
Use the applicant preferences to make the recommendations more specific.
--- APPLICANT PREFERENCES ---
{application_context}
--- RESUME ---
{resume_text}
--- JOB DESCRIPTION ---
{job_desc_text}
{OUTPUT_SCHEMA}
"""

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            response_format={"type": "json_object"},
            temperature=DEFAULT_TEMPERATURE,
        )
        return json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        raise HTTPException(status_code=502, detail="The analysis returned an unexpected format. Please try again.")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=502, detail="Could not complete the analysis right now. Please try again in a minute.")


@app.post("/api/analyze")
async def analyze(payload: AnalyzeRequest):
    if not payload.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text is empty.")
    if not payload.job_desc_text.strip():
        raise HTTPException(status_code=400, detail="Job description text is empty.")

    result = analyze_resume_ats(
        resume_text=payload.resume_text,
        job_desc_text=payload.job_desc_text,
        target_role=payload.target_role,
        experience_level=payload.experience_level,
        focus_area=payload.focus_area,
    )
    return result


@app.get("/api/health")
async def health():
    return {"status": "ok", "groq_configured": bool(os.getenv("GROQ_API_KEY"))}
